import docx
import base64
import os
from docx.shared import Inches, Pt, RGBColor
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
import requests
import io
from PIL import Image
import numpy as np
from functions.Sharepoint.get_sharepoint_columns import get_sharepoint_access_headers_through_client_id
from functions.Image_api import resize_and_autoorient

def create_word_document(js, doc=None):
    
    if not doc: doc = docx.Document()
    elif isinstance(doc, io.BytesIO):
        doc = docx.Document(doc)
    elif isinstance(doc,str):
        doc = docx.Document(io.BytesIO(base64.b64decode(doc)))
    
    for item in js:

        if "Paragraphs" in item.keys():
            for para in item['Paragraphs']:
                add_paragraph(doc,para['runs'])
        elif "Table" in item.keys():
            create_word_table_from_json(doc,item['Table'])
        elif "Footer" in item.keys():
            create_footer(doc,item['Footer'])
        elif "Header" in item.keys():
            create_header(doc,item['Header'])
        elif "Page_break" in item.keys():
            doc.add_page_break()
    file = io.BytesIO()
    doc.save(file)
    file.seek(0)
    return {"$content":base64.b64encode(file.getvalue()).decode('utf-8'),"$content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}


def add_paragraph(doc, js=[{'text':"Hello",
                            'params':{
                                'hyperlink':{'url':'http://www.aendraes.com'}
                            }}]):
    p = doc.add_paragraph()
    fill_paragraph(doc,p,js)


def fill_paragraph(doc, paragraph, paragraph_content):
    for run in paragraph_content:
        run0 = paragraph.add_run()
        custom_style = "custom_style"
        while custom_style in doc.styles:
            custom_style = custom_style + '1'
        rstyle = doc.styles.add_style(custom_style, WD_STYLE_TYPE.CHARACTER)
        run0.style=rstyle
        if 'text' in run.keys(): 
            run0.text = run['text']
        if 'params' in run.keys():
            params = run['params']
            if "keep_with_next" in params.keys():
                paragraph.keep_with_next=params["keep_with_next"]
            if "border" in params.keys():
                border = params["border"]
                if "color" in border.keys(): color = border["color"].replace('#','')
                else: color="auto"
                
                if "area" in border.keys(): area = border["area"]
                else: area = "bottom"
                if "border-type" in border.keys(): border_type = border["border-type"]
                else: border_type = "single"
                if "size" in border.keys(): size = border["size"]
                else: size=1
                if "space" in border.keys(): space = border["space"]
                else: space=1
                line = f"""
                <w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                            <w:{area} w:val="{border_type}" w:sz="{size}" w:space="{space}" w:color="{color}"/>
                            </w:pBdr>
            """
                pBdr = parse_xml(line)
                pPr = paragraph._element.get_or_add_pPr()
                pPr.append(pBdr)
            if "hyperlink" in params.keys():
                if 'url' in params['hyperlink'].keys():
                    add_hyperlink(paragraph,run0, params['hyperlink']['url'])
            if "font" in params.keys():
                fontjs = params['font']

                font=rstyle.font
                if "size" in fontjs.keys():
                    font.size = Pt(int(fontjs['size']))
                if "name" in fontjs.keys():
                    font.name = fontjs['name']
                if "bold" in fontjs.keys():
                    font.bold = fontjs['bold']
                if "italic" in fontjs.keys():
                    font.italic = fontjs['italic']
                if "underline" in fontjs.keys():
                    font.underline = fontjs['underline']
            if "color" in params.keys():
                font.color.rgb = RGBColor(*params['color'])
            if "alignment" in params.keys():
                if params["alignment"]=='center': paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif params["alignment"]=='left': paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif params['alignment']=='right': paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif params['alignment']=='justify': paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                else: paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif "image" in run.keys():
            if run["image"]["content"].startswith('http'):
                headers = {}
                if 'headers' in run['image'].keys(): 
                    headers=run['image']['headers']
                elif "sharepoint" in run["image"]["content"]:
                    headers = get_sharepoint_access_headers_through_client_id()
                rs = requests.get(run["image"]["content"],headers=headers)
                img = base64.b64encode(rs.content).decode('utf-8')
                run['image']['content'] = img
            if "content" in run["image"].keys():
                img = io.BytesIO(base64.b64decode(run["image"]['content']))
            else: img = io.BytesIO(base64.b64decode(run["image"]['$content']))
            img.seek(0)
            height=0.5
            width=1.5
            if "size" in run["image"].keys():
                width,height = run["image"]["size"]
            run0.add_picture(img, width=Inches(width))#, height=Inches(height))
        elif "checkbox" in run.keys():
            checked = False
            size = 20
            if "checked" in run['checkbox'].keys(): checked = run["checkbox"]["checked"]
            if "size" in run["checkbox"].keys(): size = run["checkbox"]["size"]
            add_checkbox(paragraph,checked=checked,size=size)

def create_footer(doc, js):
    section = doc.sections[0]
    footer = section.footer
    if 'paragraphs' in js.keys():
        paragraphs = js['paragraphs']
        
        for para in paragraphs:
            p = footer.add_paragraph()
            if "runs" in para.keys():
                fill_paragraph(doc, p,paragraph_content=para["runs"])
    return doc

def create_word_table_from_json(doc, js):
    params = {}
    if "params" in js.keys(): params = js['params']
    js = js['items']
    if not any(js): return doc
    ## Handling of input parameters
    header_row = None
    image_columns=[]
    col_width=[]
    hyperlink_columns = []
    pstyle = 'Normal'
    image_size=[140,140]
    columns = list(js[0].keys())
    #print(columns)
    table_style,preset,paragraph, image_properties, paragraph_style, rename_columns = None,None,None, None, None, None
    if params:

        if "column_widths" in params.keys(): col_width = params['column_widths']
        if "preset" in params.keys(): preset = params['preset']
        if "paragraph" in params.keys():paragraph = params['paragraph']
        if "table_style" in params.keys(): table_style = params['table_style']
        if "image_properties" in params.keys(): image_properties = params['image_properties']
        if "columns" in params.keys(): columns = params['columns']
        if "header_row" in params.keys(): header_row = params['header_row']
        if 'hyperlink-columns' in params.keys(): hyperlink_columns = params['hyperlink-columns']
        
    if image_properties:
        if "columns" in image_properties.keys():image_columns = image_properties['columns']
        if "image_size" in image_properties.keys(): image_size = image_properties['image_size']
    
        
    if paragraph:
        if "paragraph_style" in paragraph.keys(): p_style = paragraph['paragraph_style']
        else: 
            styles = doc.styles
            randstyle = 'Customtablestyle'+str(np.random.randint(1, 1000001))
            if randstyle in styles: randstyle = 'Customtablestyle'+str(np.random.randint(1000001, 2000001))
            pstyle = randstyle
            style = styles.add_style(randstyle, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles['Normal']
            if "font" in paragraph.keys():
                fontjs = paragraph['font']
                font=style.font
                if "size" in fontjs.keys():
                    font.size = Pt(int(fontjs['size']))
                if "name" in fontjs.keys():
                    font.name = fontjs['name']
                if "bold" in fontjs.keys():
                    font.bold = fontjs['bold']
                if "italic" in fontjs.keys():
                    font.italic = fontjs['italic']
                if "underline" in fontjs.keys():
                    font.underline = fontjs['underline']
                if "name" in fontjs.keys():
                    font.name=fontjs["name"]
            if "color" in paragraph.keys():
                R, G , B = paragraph['color']
                font.color.rgb = RGBColor(R, G, B)
            pstyle = randstyle
    for i, item in enumerate(js):
        if "rename_columns" in params.keys():
                for old,new in zip(*params["rename_columns"]):
                    if old not in item.keys():
                        item[old] = ""
                    item[new] = item[old]
                columns = params["rename_columns"][1]
    #if image_columns:
        #print([item[image_column] for image_column in image_columns for item in js])
    #    if not any([item[image_column] for image_column in image_columns for item in js]): return doc
    table = doc.add_table(rows=1, cols = len(columns))
    if table_style:
        table.style=table_style
        
    for i, item in enumerate(js):
        if not any(item):continue
        ## Sätt rubriker till tabellen
        if "header_row" in params.keys() and i==0:
            header_row = params["header_row"]
            if "height" in header_row.keys():
                height = header_row["height"]
            else: height=0.2
            if "background-color" in header_row.keys():
                bcolor = header_row["background-color"]
                bcolor = '{:02x}{:02x}{:02x}'.format(*bcolor)
            else: bcolor="22E42F"
            if "font-size" in header_row.keys(): font_size=header_row["font-size"]
            else: font_size=12
            if "bold" in header_row.keys(): bold=header_row["bold"]
            else: bold = True
            if "italic" in header_row.keys(): italic=header_row["italic"]
            else: italic = False
            if "underline" in header_row.keys(): underline=header_row["underline"]
            else: underline = False
            
            
            header_row = table.rows[0].cells
            for i in range(len(columns)):
                p = header_row[i].paragraphs[0]
                p.paragraph_format.keep_with_next=True
                p.text = columns[i]
                p.runs[0].bold=bold
                p.runs[0].italic=italic
                p.runs[0].underline=underline
                p.runs[0].font.size = Pt(font_size)
                fillxml = f"""<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:color="auto" w:fill="{bcolor}" />"""
                
                header_row[i]._tc.get_or_add_tcPr().append(parse_xml(fillxml))
            table.rows[0].height=Inches(height)
            

        row = table.add_row().cells
        
        for i, column in enumerate(columns):
            if column not in item.keys(): item[column] = " "
            if column in image_columns:
                images = item[column]
                
                if not any(images):break
                #print(item['Images'])
                if any(images):
                    for image in images: 
                        if image.startswith('http'):
                            imagefile = requests.get(image)
                            if imagefile.status_code != 200: continue
                            imagefile = imagefile.content
                            imagefile = base64.b64encode(imagefile).decode('utf-8')
                        else:
                            imagefile = image
                        imagefile.replace('\n','')
                        #while len(imagefile)%4 != 0: imagefile = imagefile + "="
                        file = io.BytesIO(base64.b64decode(imagefile))
                        file.seek(0)
                        file = resize_and_autoorient(file,image_size[0])
                        p=row[i].paragraphs[0]
                        p.style.paragraph_format.keep_with_next=True
                        run = p.add_run()
                        picture = run.add_picture(file)
                    
                continue
            elif column == []: continue
            
            p = row[i].paragraphs[0]
            
            p.style = pstyle
            p.paragraph_format.keep_with_next=True
            if column in hyperlink_columns: 
                add_hyperlink(p,str(item[column]),str(item[column]))
            else:
                p.text = str(item[column])
            
    if col_width:
        for i, column in enumerate(table.columns):
            width = col_width[i]
            for cell in column.cells:
                cell.width = Inches(width)
                
    return doc

def create_header(doc, js):
    section = doc.sections[0]
    header = section.header
    if 'paragraphs' in js.keys():
        paragraphs = js['paragraphs']
        
        for para in paragraphs:
            p = header.add_paragraph()
            if "runs" in para.keys():
                fill_paragraph(doc, p,paragraph_content=para["runs"])
    return doc

def add_checkbox(paragraph, checked=False, size=20):
    box = "☐"
    checkvalue = 0
    if checked:
        box="☒"
        checkvalue = 1
    XML = f"""
                     <w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
                xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                xmlns:ns0="http://schemas.openxmlformats.org/markup-compatibility/2006">
            <w:sdtPr>
                <w:rPr>
                    <w:sz w:val="{size}"/>
                    <w:szCs w:val="{size}"/>
                </w:rPr>
                <w:id w:val="-134335307"/>
                <w14:checkbox>
                    <w14:checked w14:val="{checkvalue}"/>
                    <w14:checkedState w14:val="2612" w14:font="MS Gothic"/>
                    <w14:uncheckedState w14:val="2610" w14:font="MS Gothic"/>
                </w14:checkbox>
            </w:sdtPr>
            <w:sdtEndPr/>
            <w:sdtContent>
                <w:r w:rsidR="007415F5">
                    <w:rPr>
                        <w:rFonts w:ascii="MS Gothic" w:eastAsia="MS Gothic" w:hAnsi="MS Gothic" w:hint="eastAsia"/>
                        <w:sz w:val="{size}"/>
                        <w:szCs w:val="{size}"/>
                    </w:rPr>
                    <w:t>{box}</w:t>
                </w:r>
            </w:sdtContent>
        </w:sdt>
     """
    std_element = parse_xml(XML)
    paragraph._p.append(std_element)
    paragraph.add_run(" Text2")
    return None



def add_hyperlink(paragraph, run, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = run._r
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append (hyperlink)
    return hyperlink


if __name__ == '__main__':
    import json
    with open(os.path.join(os.path.dirname(__file__),'sample.json'), 'r', encoding='utf-8') as f :
        js = json.load(f)
    print(type(js['body']))
    js = js['body']
    #js = json.loads(js['body'])
    doc = create_word_document(js)
    
    with open(os.path.join(os.path.dirname(__file__),'sample.docx'), 'wb') as f :
        f.write(base64.b64decode(doc['$content']))