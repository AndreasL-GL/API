from __future__ import print_function
import mailmerge
import docx
from PIL import Image
import base64
import io
import json, os
from mailmerge import MailMerge
import Printer

import functions.Word.get_template as get_template
import functions.Word.create_json as create_json
from functions.Word.add_image_to_zipfile import add_icon_to_word_file, autoorient


def add_template_data(file, tbl):
    """ Adds the data into the template file that needs to be added with MailMerge. This is the Table without the images.

    Args:
        file (BytesIO): Document file, downloaded from sharepoint-site.
        tbl (dict): Json-dictionary contained in the request body.

    Returns:
        _type_: returns a BytesIO object to be opened by python-docx.
    """
    doc = mailmerge.MailMerge(file)
    content= tbl["content"]
    info = {k:v for k,v in tbl.items() }#if k!="content"}

    doc.merge(**info)
    doc.merge_rows('Område',content)
    
    
    
    for item in content:
        item["Område2"] = item["Område"]
    file = io.BytesIO()
    doc.write(file)
    file.seek(0)
    return file


def add_images_to_doc(file,tbl):
    """Adds all the image rows with the accompanying test per the template we used before."""
    doc = docx.Document(file)
    
    doc.add_heading('Bilder och Kommentarer', 0)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Plain Table 5'
    for item in tbl["images"]:
        row = table.add_row().cells
        Adress = item["Område"]
        Bildkommentar = item["Bildkommentar"]
        Bild = item["Bild"]
        row[0].text = Adress
        row[1].text = Bildkommentar

        img = Image.open(Bild)
        img =autoorient(img)
        img = img.resize((200,250))
        file=io.BytesIO()
        img.save(file, format='JPEG')
        file.seek(0)
        row[2].add_paragraph().add_run().add_picture(file)
    return doc

def change_icon_in_header(doc):
    """Changes the icon in the header of a word file by replacing a pre-formatted placeholder image."""
    l = get_template.Download_icon()
    file_content = io.BytesIO(l.content)
    file_content.seek(0)
    img = Image.open(file_content)
    fb = io.BytesIO()
    img.save(fb, 'PNG')
    fb.seek(0)
    docbyte = io.BytesIO()
    doc.save(docbyte)
    docbyte.seek(0)
    doc = add_icon_to_word_file(docbyte.read(), icon_file=fb.read(), imagepath="word/media/image2.png")
    doc = docx.Document(doc)
    return doc

def run_functions(js):
    """Runs all the different functions. Takes json as input, returns a json with file-content and filename."""
    title = js["Title"]
    kontrollmoment = get_template.get_fields(list_=title)    
    momentstr = ""
    with open('001file.json', "w", encoding='utf-8') as f:
        json.dump(kontrollmoment,f, indent=4, ensure_ascii=False)
    js["Kontrollmoment"] = kontrollmoment
    for i,item in enumerate(js['Items']['value']):
        momentstr = ""
        for jitem in item['Kontrollmoment']:
            momentstr = momentstr+'- '+jitem['Value']+'\n'
        js['Items']['value'][i]['Kontrollmoment'] = momentstr
    js['Kontrollmoment'] = [item for item in js['Kontrollmoment'] if item['Moment'] in momentstr]
    js = create_json.create_json_for_word_functions(js)
    file= download_template_file()
    file = add_template_data(file, js)
    doc = add_images_to_doc(file, js)
    doc = change_icon_in_header(doc)
    file_content = io.BytesIO()
    doc.save(file_content)
    doc.save('testfile.docx')
    doc.save(title+" Vecka "+js["Vecka"]+'.docx')
    file_content.seek(0)
    file_content=base64.b64encode(file_content.getvalue()).decode('utf-8')
    return {"content":file_content, "filename": title+" Vecka "+js["Vecka"]}

def mailmerge_fun(doc,js):
    print(type(doc))
    if type(doc)==str: doc = MailMerge(io.BytesIO(base64.b64decode(doc)))
    elif type(doc)==dict: 
        print(doc.keys())
        doc=MailMerge(io.BytesIO(base64.b64decode(doc["$content"])))
    
    elif isinstance(doc,io.BytesIO):
        doc.seek(0) 
        doc=MailMerge(doc.getvalue())
    elif type(doc)==bytes: doc=MailMerge(bytes)
    
    
    if isinstance(js,list): js=js[0]
    doc.merge(**js)
    file = io.BytesIO()
    doc.write(file)
    file.seek(0)
    return {"$content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "$content":base64.b64encode(file.getvalue()).decode('utf-8')}

def download_template_file():
    """Downloads the word template for this program."""
    l = get_template.Download_template("","word_template.docx")
    file_content = io.BytesIO(l.content)
    file_content.seek(0)
    return file_content
if __name__ == '__main__':
    
    #with open('items.json', 'r', encoding="utf-8") as f:
    #    js = json.load(f, )
   # run_functions(js)
    with open(os.path.join(os.path.dirname(__file__),'mailmergetest.json'), encoding='utf-8') as f:
        js = json.load(f)['body']
    doc = mailmerge_fun(js['Document'],js['Items'])
    with open(os.path.join(os.path.dirname(__file__),'mailmergetest2.docx'),mode='wb') as f:
        f.write(base64.b64decode(doc["$content"]))